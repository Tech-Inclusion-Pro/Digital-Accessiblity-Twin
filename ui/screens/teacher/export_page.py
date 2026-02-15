"""Teacher export page — AI-powered DOCX student report builder."""

import asyncio
import json
from datetime import datetime, timezone

from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton,
    QScrollArea, QComboBox, QFrame, QDialog, QCheckBox, QTextEdit,
    QFileDialog, QMessageBox,
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QTimer

from config.settings import get_colors, APP_SETTINGS
from models.student_profile import StudentProfile
from models.support import SupportEntry
from models.tracking import TrackingLog
from models.consultation_log import ConsultationLog
from models.document import Document
from models.evaluation import TwinEvaluation
from models.insight_log import InsightLog
from ui.components.empty_state import EmptyState
from ui.components.mic_button import MicButton


# ------------------------------------------------------------------ workers

class _ReportStreamWorker(QThread):
    """Runs the async AI report generation in a background thread."""

    chunk_received = pyqtSignal(str)
    finished_signal = pyqtSignal()
    error_signal = pyqtSignal(str)

    def __init__(self, backend_manager, system_prompt, parent=None):
        super().__init__(parent)
        self.bm = backend_manager
        self.system_prompt = system_prompt

    def run(self):
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            loop.run_until_complete(self._stream())
        except Exception as e:
            self.error_signal.emit(str(e))
        finally:
            loop.close()
            self.finished_signal.emit()

    async def _stream(self):
        async for chunk in self.bm.generate_response(
            "Please generate the student accessibility report now.",
            system_prompt=self.system_prompt,
            conversation_history=[],
        ):
            self.chunk_received.emit(chunk)


# ------------------------------------------------------------------ dialog

class ReportConfigDialog(QDialog):
    """Popup dialog for configuring which sections to include and AI guidance."""

    def __init__(self, student_name: str, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"Report Configuration \u2014 {student_name}")
        self.setMinimumSize(540, 520)
        c = get_colors()
        self.setStyleSheet(f"QDialog {{ background-color: {c['dark_bg']}; }}")

        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 20, 24, 20)
        layout.setSpacing(14)

        # Header
        header = QLabel("Configure Report Sections")
        header.setStyleSheet(
            f"font-size: 18px; font-weight: bold; color: {c['text']};"
        )
        layout.addWidget(header)

        desc = QLabel(
            "Toggle the sections you want included in the report. "
            "Use the guidance area to tell the AI how to shape the report."
        )
        desc.setWordWrap(True)
        desc.setStyleSheet(f"font-size: 13px; color: {c['text_muted']};")
        layout.addWidget(desc)

        # Separator
        sep = QFrame()
        sep.setFrameShape(QFrame.Shape.HLine)
        sep.setStyleSheet(f"color: {c['dark_border']};")
        layout.addWidget(sep)

        # Section toggles
        sections_label = QLabel("Include in Report:")
        sections_label.setStyleSheet(
            f"font-size: 14px; font-weight: bold; color: {c['text']};"
        )
        layout.addWidget(sections_label)

        cb_style = f"""
            QCheckBox {{
                color: {c['text']}; font-size: 13px; spacing: 8px;
                min-height: {APP_SETTINGS['touch_target_min']}px;
            }}
            QCheckBox::indicator {{
                width: 20px; height: 20px;
            }}
        """

        self._profile_cb = QCheckBox("Profile Summary (strengths, hopes, stakeholders)")
        self._profile_cb.setChecked(True)
        self._profile_cb.setStyleSheet(cb_style)
        layout.addWidget(self._profile_cb)

        self._supports_cb = QCheckBox("Support Entries (categories, effectiveness ratings)")
        self._supports_cb.setChecked(True)
        self._supports_cb.setStyleSheet(cb_style)
        layout.addWidget(self._supports_cb)

        self._tracking_cb = QCheckBox("Tracking Logs (implementation & outcome notes)")
        self._tracking_cb.setChecked(True)
        self._tracking_cb.setStyleSheet(cb_style)
        layout.addWidget(self._tracking_cb)

        self._consultations_cb = QCheckBox("Consultation History (coach conversations)")
        self._consultations_cb.setChecked(True)
        self._consultations_cb.setStyleSheet(cb_style)
        layout.addWidget(self._consultations_cb)

        self._insights_cb = QCheckBox("AI Insights (generated analysis reports)")
        self._insights_cb.setChecked(True)
        self._insights_cb.setStyleSheet(cb_style)
        layout.addWidget(self._insights_cb)

        self._chats_cb = QCheckBox("Insight Chat Conversations")
        self._chats_cb.setChecked(False)
        self._chats_cb.setStyleSheet(cb_style)
        layout.addWidget(self._chats_cb)

        # Separator
        sep2 = QFrame()
        sep2.setFrameShape(QFrame.Shape.HLine)
        sep2.setStyleSheet(f"color: {c['dark_border']};")
        layout.addWidget(sep2)

        # AI Guidance
        guidance_label = QLabel("AI Report Guidance:")
        guidance_label.setStyleSheet(
            f"font-size: 14px; font-weight: bold; color: {c['text']};"
        )
        layout.addWidget(guidance_label)

        guidance_desc = QLabel(
            "Describe how you would like the AI to write this report. "
            "For example, what to focus on, the intended audience, or the tone."
        )
        guidance_desc.setWordWrap(True)
        guidance_desc.setStyleSheet(f"font-size: 12px; color: {c['text_muted']};")
        layout.addWidget(guidance_desc)

        self._guidance = QTextEdit()
        self._guidance.setPlaceholderText(
            "e.g., Focus on reading and writing supports. Highlight areas of "
            "growth this semester. This report is for the parent-teacher "
            "conference. Keep the tone positive and strengths-based."
        )
        self._guidance.setAccessibleName("AI report guidance")
        self._guidance.setFixedHeight(120)
        self._guidance.setStyleSheet(f"""
            QTextEdit {{
                background: {c['dark_input']}; color: {c['text']};
                border: 1px solid {c['dark_border']}; border-radius: 8px;
                padding: 8px; font-size: 13px;
            }}
        """)
        guidance_row = QHBoxLayout()
        guidance_row.addWidget(self._guidance, stretch=1)
        guidance_row.addWidget(MicButton(target=self._guidance), alignment=Qt.AlignmentFlag.AlignTop)
        layout.addLayout(guidance_row)

        # Buttons
        btn_row = QHBoxLayout()
        btn_row.setSpacing(10)
        btn_row.addStretch()

        cancel_btn = QPushButton("Cancel")
        cancel_btn.setAccessibleName("Cancel report configuration")
        cancel_btn.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
        cancel_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        cancel_btn.setFixedHeight(APP_SETTINGS["touch_target_min"])
        cancel_btn.setStyleSheet(f"""
            QPushButton {{
                background: {c['dark_input']}; color: {c['text']};
                border: 1px solid {c['dark_border']}; border-radius: 8px;
                padding: 0 24px; font-size: 13px;
            }}
        """)
        cancel_btn.clicked.connect(self.reject)
        btn_row.addWidget(cancel_btn)

        generate_btn = QPushButton("Generate Report")
        generate_btn.setAccessibleName("Generate the report with selected options")
        generate_btn.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
        generate_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        generate_btn.setFixedHeight(APP_SETTINGS["touch_target_min"])
        generate_btn.setStyleSheet(f"""
            QPushButton {{
                background: {c['primary']}; color: white;
                border: none; border-radius: 8px;
                padding: 0 24px; font-weight: bold; font-size: 13px;
            }}
        """)
        generate_btn.clicked.connect(self.accept)
        btn_row.addWidget(generate_btn)

        layout.addLayout(btn_row)

    @property
    def config(self) -> dict:
        return {
            "profile": self._profile_cb.isChecked(),
            "supports": self._supports_cb.isChecked(),
            "tracking": self._tracking_cb.isChecked(),
            "consultations": self._consultations_cb.isChecked(),
            "insights": self._insights_cb.isChecked(),
            "chats": self._chats_cb.isChecked(),
            "guidance": self._guidance.toPlainText().strip(),
        }


# ------------------------------------------------------------------ page

class TeacherExportPage(QWidget):
    """AI-powered student report builder with DOCX export."""

    def __init__(self, db_manager, auth_manager, backend_manager=None, parent=None):
        super().__init__(parent)
        self.db = db_manager
        self.auth = auth_manager
        self.backend_manager = backend_manager
        self._worker = None
        self._student_data: list = []  # [(profile_id, name)]
        self._report_text = ""
        self._current_student_name = ""
        self._build_ui()

    def _build_ui(self):
        c = get_colors()
        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 20, 24, 20)
        layout.setSpacing(12)

        # Header
        header = QLabel("Export Student Report")
        header.setStyleSheet(
            f"font-size: 22px; font-weight: bold; color: {c['text']};"
        )
        header.setAccessibleName("Export Student Report")
        layout.addWidget(header)

        desc = QLabel(
            "Generate an AI-powered accessibility report for a student. "
            "Choose which sections to include and provide guidance to shape "
            "the report. Export as a timestamped Word document."
        )
        desc.setWordWrap(True)
        desc.setStyleSheet(f"font-size: 13px; color: {c['text_muted']};")
        layout.addWidget(desc)

        # Student selector row
        selector_row = QHBoxLayout()
        selector_row.setSpacing(10)

        selector_label = QLabel("Student:")
        selector_label.setStyleSheet(f"font-size: 14px; color: {c['text']};")
        selector_row.addWidget(selector_label)

        self._combo = QComboBox()
        self._combo.setAccessibleName("Select student for report")
        self._combo.setFixedHeight(APP_SETTINGS["touch_target_min"])
        self._combo.setMinimumWidth(250)
        self._combo.setStyleSheet(f"""
            QComboBox {{
                background: {c['dark_input']}; color: {c['text']};
                border: 1px solid {c['dark_border']}; border-radius: 8px;
                padding: 0 12px; font-size: 13px;
            }}
        """)
        selector_row.addWidget(self._combo)

        self._create_btn = QPushButton("Create Report")
        self._create_btn.setAccessibleName("Open report configuration dialog")
        self._create_btn.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
        self._create_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self._create_btn.setFixedHeight(APP_SETTINGS["touch_target_min"])
        self._create_btn.setStyleSheet(f"""
            QPushButton {{
                background: {c['primary']}; color: white;
                border: none; border-radius: 8px; padding: 0 20px;
                font-weight: bold; font-size: 13px;
            }}
            QPushButton:disabled {{
                background: {c['dark_border']}; color: {c['text_muted']};
            }}
        """)
        self._create_btn.clicked.connect(self._open_config_dialog)
        selector_row.addWidget(self._create_btn)

        selector_row.addStretch()
        layout.addLayout(selector_row)

        # Status label
        self._status_label = QLabel("")
        self._status_label.setStyleSheet(
            f"font-size: 12px; color: {c['text_muted']};"
        )
        layout.addWidget(self._status_label)

        # Separator
        sep = QFrame()
        sep.setFrameShape(QFrame.Shape.HLine)
        sep.setStyleSheet(f"color: {c['dark_border']};")
        layout.addWidget(sep)

        # Empty state
        self._empty = EmptyState(
            icon_text="\u25A1",
            message="No students found. Import student twins first.",
            action_label="Go to Students",
        )
        self._empty.setVisible(False)
        layout.addWidget(self._empty)

        # Report preview area
        self._output_scroll = QScrollArea()
        self._output_scroll.setWidgetResizable(True)
        self._output_scroll.setStyleSheet(
            f"QScrollArea {{ border: 1px solid {c['dark_border']}; "
            f"border-radius: 8px; background: {c['dark_card']}; }}"
        )
        self._output_scroll.setAccessibleName("Report preview area")

        self._output_label = QLabel(
            "Select a student and click 'Create Report' to get started."
        )
        self._output_label.setWordWrap(True)
        self._output_label.setAlignment(Qt.AlignmentFlag.AlignTop)
        self._output_label.setStyleSheet(
            f"font-size: 13px; color: {c['text_muted']}; padding: 16px;"
        )
        self._output_label.setTextInteractionFlags(
            Qt.TextInteractionFlag.TextSelectableByMouse
            | Qt.TextInteractionFlag.TextSelectableByKeyboard
        )
        self._output_scroll.setWidget(self._output_label)
        layout.addWidget(self._output_scroll, stretch=1)

        # Export button row
        export_row = QHBoxLayout()
        export_row.setSpacing(10)

        self._export_btn = QPushButton("Export as Word (.docx)")
        self._export_btn.setAccessibleName("Export report as Word document")
        self._export_btn.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
        self._export_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self._export_btn.setFixedHeight(APP_SETTINGS["touch_target_min"])
        self._export_btn.setEnabled(False)
        self._export_btn.setStyleSheet(f"""
            QPushButton {{
                background: {c['tertiary']}; color: white;
                border: none; border-radius: 8px;
                padding: 0 24px; font-weight: bold; font-size: 13px;
            }}
            QPushButton:disabled {{
                background: {c['dark_border']}; color: {c['text_muted']};
            }}
        """)
        self._export_btn.clicked.connect(self._export_docx)
        export_row.addWidget(self._export_btn)

        export_row.addStretch()
        layout.addLayout(export_row)

    # --------------------------------------------------------- dialog + generation

    def _open_config_dialog(self):
        """Open the report configuration popup dialog."""
        index = self._combo.currentIndex()
        if index < 0 or index >= len(self._student_data):
            return

        profile_id, name = self._student_data[index]
        self._current_student_name = name

        if not self.backend_manager or self.backend_manager._client is None:
            self._output_label.setText(
                "No AI backend configured. Please set up an AI provider in "
                "AI Settings before generating reports."
            )
            return

        dlg = ReportConfigDialog(name, self)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return

        config = dlg.config
        self._generate_report(profile_id, name, config)

    def _generate_report(self, profile_id: int, name: str, config: dict):
        """Gather data, build the AI prompt, and start report generation."""
        self._create_btn.setEnabled(False)
        self._export_btn.setEnabled(False)
        self._output_label.setText("Generating report...")
        self._report_text = ""

        # Gather data from database
        session = self.db.get_session()
        try:
            profile = session.query(StudentProfile).get(profile_id)
            if not profile:
                self._output_label.setText("Student profile not found.")
                self._create_btn.setEnabled(True)
                return

            data_sections = []
            first_name = name.split()[0] if name else "Student"
            now = datetime.now(timezone.utc).strftime("%B %d, %Y")

            # Profile
            if config["profile"]:
                items = []
                for field, label in [
                    ("strengths", "Strengths"),
                    ("history", "History"),
                    ("hopes", "Hopes & Goals"),
                    ("stakeholders", "Stakeholders"),
                ]:
                    raw = getattr(profile, field)
                    if raw:
                        texts = []
                        for item in raw:
                            if isinstance(item, dict):
                                texts.append(item.get("text", str(item)))
                            else:
                                texts.append(str(item))
                        items.append(f"{label}: {'; '.join(texts)}")
                if items:
                    data_sections.append(
                        "PROFILE SUMMARY:\n" + "\n".join(items)
                    )

            # Supports
            if config["supports"]:
                supports = session.query(SupportEntry).filter(
                    SupportEntry.profile_id == profile_id
                ).all()
                if supports:
                    lines = []
                    for s in supports:
                        rating = f" (effectiveness: {s.effectiveness_rating}/5)" if s.effectiveness_rating else ""
                        lines.append(
                            f"  [{s.category}] {s.description} — {s.status}{rating}"
                        )
                    data_sections.append(
                        f"SUPPORT ENTRIES ({len(supports)}):\n" + "\n".join(lines)
                    )

            # Tracking logs
            if config["tracking"]:
                logs = session.query(TrackingLog).filter(
                    TrackingLog.profile_id == profile_id,
                ).order_by(TrackingLog.created_at.desc()).limit(30).all()
                if logs:
                    lines = []
                    for lg in logs:
                        ts = lg.created_at.strftime("%Y-%m-%d") if lg.created_at else "unknown"
                        impl = (lg.implementation_notes or "")[:200]
                        outcome = (lg.outcome_notes or "")[:200]
                        lines.append(f"  {ts}: Impl: {impl} | Outcome: {outcome}")
                    data_sections.append(
                        f"TRACKING LOGS ({len(logs)}):\n" + "\n".join(lines)
                    )

            # Consultations
            if config["consultations"]:
                consults = (
                    session.query(ConsultationLog)
                    .filter(ConsultationLog.profile_id == profile_id)
                    .order_by(ConsultationLog.created_at.desc())
                    .limit(10)
                    .all()
                )
                if consults:
                    lines = []
                    for i, cl in enumerate(consults, 1):
                        ts = cl.created_at.strftime("%Y-%m-%d") if cl.created_at else "unknown"
                        summary = cl.summary or "(no summary)"
                        lines.append(
                            f"  Consultation {i} ({ts}): {summary} "
                            f"({cl.message_count} messages)"
                        )
                    data_sections.append(
                        f"CONSULTATION HISTORY ({len(consults)}):\n"
                        + "\n".join(lines)
                    )

            # AI Insights
            if config["insights"]:
                user = self.auth.get_current_user()
                insights = (
                    session.query(InsightLog)
                    .filter(
                        InsightLog.profile_id == profile_id,
                        InsightLog.role == "teacher",
                    )
                    .order_by(InsightLog.created_at.desc())
                    .limit(5)
                    .all()
                )
                if insights:
                    lines = []
                    for ins in insights:
                        ts = ins.created_at.strftime("%Y-%m-%d") if ins.created_at else "unknown"
                        content_preview = ins.content[:500]
                        lines.append(f"  --- Insight ({ts}) ---\n  {content_preview}")
                        if config["chats"] and ins.conversation:
                            for msg in ins.conversation[:6]:
                                role = "Teacher" if msg.get("role") == "user" else "Assistant"
                                lines.append(
                                    f"    {role}: {msg.get('content', '')[:200]}"
                                )
                    data_sections.append(
                        f"AI INSIGHTS ({len(insights)}):\n" + "\n".join(lines)
                    )
        finally:
            session.close()

        # Build the AI prompt
        guidance = config["guidance"] or "Write a comprehensive overview of the student's accessibility supports and progress."

        system_prompt = (
            "You are writing a professional accessibility support report for a teacher "
            "about a student. This report may be shared with administrators, parents, "
            "or support teams.\n\n"
            "IMPORTANT PRIVACY RULES:\n"
            "- Use the student's first name only\n"
            "- Never mention specific diagnoses or medical conditions\n"
            "- Focus on supports, strengths, and progress\n"
            "- Keep language strengths-based and respectful\n\n"
            f"Student first name: {first_name}\n"
            f"Report date: {now}\n\n"
            f"TEACHER'S GUIDANCE:\n{guidance}\n\n"
            "STUDENT DATA:\n" + "\n\n".join(data_sections) + "\n\n"
            "Write a well-structured professional report using markdown formatting:\n"
            "- Use ## for section headings\n"
            "- Use ### for sub-headings\n"
            "- Use bullet points (- ) where appropriate\n"
            "- Use **bold** for emphasis\n"
            "- Begin with a brief executive summary\n"
            "- Focus on actionable insights and progress\n"
            "- End with recommendations for next steps"
        )

        # Start streaming
        self._output_label.setText("")
        self._worker = _ReportStreamWorker(self.backend_manager, system_prompt)
        self._worker.chunk_received.connect(self._on_chunk)
        self._worker.finished_signal.connect(self._on_stream_done)
        self._worker.error_signal.connect(self._on_stream_error)
        self._worker.start()

    def _on_chunk(self, chunk: str):
        current = self._output_label.text()
        self._output_label.setText(current + chunk)
        self._report_text += chunk
        QTimer.singleShot(10, self._scroll_to_bottom)

    def _on_stream_done(self):
        self._create_btn.setEnabled(True)
        self._export_btn.setEnabled(True)
        self._status_label.setText("Report generated. Click 'Export as Word' to save.")

    def _on_stream_error(self, error_msg: str):
        self._output_label.setText(f"Error generating report: {error_msg}")
        self._create_btn.setEnabled(True)
        self._status_label.setText("")

    def _scroll_to_bottom(self):
        vbar = self._output_scroll.verticalScrollBar()
        vbar.setValue(vbar.maximum())

    # --------------------------------------------------------- DOCX export

    def _export_docx(self):
        if not self._report_text.strip():
            QMessageBox.warning(self, "No Report", "Generate a report first.")
            return

        name_slug = self._current_student_name.replace(" ", "_")
        ts = datetime.now().strftime("%Y%m%d_%H%M")
        default_name = f"{name_slug}_report_{ts}.docx"

        path, _ = QFileDialog.getSaveFileName(
            self, "Save Report as Word Document",
            default_name,
            "Word Documents (*.docx)",
        )
        if not path:
            return

        try:
            from docx import Document
            from docx.shared import Pt, RGBColor

            doc = Document()

            # Title
            title = doc.add_heading("Student Accessibility Report", level=0)
            title.runs[0].font.color.rgb = RGBColor(0x6F, 0x2F, 0xA6)

            first_name = self._current_student_name.split()[0] if self._current_student_name else "Student"
            doc.add_paragraph(f"Student: {first_name}")
            doc.add_paragraph(
                f"Report Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
            )
            doc.add_paragraph("")

            # Parse markdown-ish AI output into DOCX
            self._markdown_to_docx(doc, self._report_text)

            doc.save(path)
            QMessageBox.information(
                self, "Exported",
                f"Report saved to:\n{path}"
            )
        except ImportError:
            QMessageBox.critical(
                self, "Missing Library",
                "The python-docx library is required for Word export.\n"
                "Install it with: pip install python-docx"
            )
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Export failed: {e}")

    @staticmethod
    def _markdown_to_docx(doc, text: str):
        """Convert simple markdown text to python-docx elements."""
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith("**") and stripped.endswith("**"):
                p = doc.add_paragraph("")
                p.add_run(stripped[2:-2]).bold = True
            else:
                # Handle inline bold: split on **
                p = doc.add_paragraph("")
                parts = stripped.split("**")
                for i, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        if i % 2 == 1:  # odd indices are bold
                            run.bold = True

    # --------------------------------------------------------- refresh

    def refresh_data(self):
        """Reload teacher's students."""
        user = self.auth.get_current_user()
        if not user:
            return

        session = self.db.get_session()
        try:
            import_docs = session.query(Document).filter(
                Document.teacher_user_id == user.id,
                Document.purpose_description == "twin_import",
            ).all()

            profile_ids = set()
            for doc in import_docs:
                evals = session.query(TwinEvaluation).filter(
                    TwinEvaluation.document_id == doc.id
                ).all()
                for ev in evals:
                    profile_ids.add(ev.student_profile_id)

            self._student_data = []
            if profile_ids:
                profiles = session.query(StudentProfile).filter(
                    StudentProfile.id.in_(profile_ids)
                ).all()
                for p in profiles:
                    self._student_data.append((p.id, p.name))
        finally:
            session.close()

        # Populate combo
        self._combo.blockSignals(True)
        self._combo.clear()
        for pid, name in self._student_data:
            self._combo.addItem(name)
        self._combo.blockSignals(False)

        if self._student_data:
            self._empty.setVisible(False)
            self._output_scroll.setVisible(True)
        else:
            self._empty.setVisible(True)
            self._output_scroll.setVisible(False)
