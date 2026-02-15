"""Student export page — JSON, Excel, and DOCX export of all twin data."""

import json
from datetime import datetime, timezone

from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton,
    QTextEdit, QFileDialog, QMessageBox,
)
from PyQt6.QtCore import Qt

from config.settings import get_colors, APP_SETTINGS
from models.student_profile import StudentProfile
from models.support import SupportEntry
from models.tracking import TrackingLog
from models.insight_log import InsightLog
from ui.components.empty_state import EmptyState


class StudentExportPage(QWidget):
    """Export accessibility twin as JSON, Excel, or DOCX."""

    def __init__(self, db_manager, auth_manager, parent=None):
        super().__init__(parent)
        self.db = db_manager
        self.auth = auth_manager
        self._profile = None
        self._build_ui()

    def _build_ui(self):
        c = get_colors()
        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 20, 24, 20)
        layout.setSpacing(16)

        header = QLabel("Export Accessibility Twin")
        header.setStyleSheet(
            f"font-size: 22px; font-weight: bold; color: {c['text']};"
        )
        header.setAccessibleName("Export Accessibility Twin")
        layout.addWidget(header)

        desc = QLabel(
            "Export your complete accessibility twin data including your profile, "
            "supports, tracking logs, and AI insights. Take your data with you "
            "in JSON, Excel, or Word format."
        )
        desc.setStyleSheet(f"font-size: 14px; color: {c['text_muted']};")
        desc.setWordWrap(True)
        layout.addWidget(desc)

        # Empty state
        self._empty = EmptyState(
            icon_text="\u25A1",
            message="No profile to export. Create your profile first.",
        )
        self._empty.setVisible(False)
        layout.addWidget(self._empty)

        # Export content container
        self._content = QWidget()
        content_layout = QVBoxLayout(self._content)
        content_layout.setContentsMargins(0, 0, 0, 0)
        content_layout.setSpacing(12)

        # Data summary
        self._summary_label = QLabel("")
        self._summary_label.setStyleSheet(
            f"font-size: 13px; color: {c['text_muted']};"
        )
        content_layout.addWidget(self._summary_label)

        # Preview
        preview_lbl = QLabel("Data Preview:")
        preview_lbl.setStyleSheet(
            f"font-size: 14px; font-weight: bold; color: {c['text']};"
        )
        content_layout.addWidget(preview_lbl)

        self._preview = QTextEdit()
        self._preview.setReadOnly(True)
        self._preview.setAccessibleName("Export data preview")
        self._preview.setStyleSheet(f"""
            QTextEdit {{
                background: {c['dark_card']}; color: {c['text']};
                border: 1px solid {c['dark_border']}; border-radius: 8px;
                font-family: monospace; font-size: 12px;
                padding: 8px;
            }}
        """)
        content_layout.addWidget(self._preview, stretch=1)

        # Export buttons row
        btn_row = QHBoxLayout()
        btn_row.setSpacing(10)

        json_btn = QPushButton("Export as JSON")
        json_btn.setAccessibleName("Export as JSON file")
        json_btn.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
        json_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        json_btn.setFixedHeight(APP_SETTINGS["touch_target_min"])
        json_btn.setStyleSheet(f"""
            QPushButton {{
                background: {c['primary']}; color: white;
                border: none; border-radius: 8px;
                padding: 0 24px; font-weight: bold; font-size: 13px;
            }}
        """)
        json_btn.clicked.connect(self._export_json)
        btn_row.addWidget(json_btn)

        excel_btn = QPushButton("Export as Excel")
        excel_btn.setAccessibleName("Export as Excel file")
        excel_btn.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
        excel_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        excel_btn.setFixedHeight(APP_SETTINGS["touch_target_min"])
        excel_btn.setStyleSheet(f"""
            QPushButton {{
                background: {c['secondary']}; color: white;
                border: none; border-radius: 8px;
                padding: 0 24px; font-weight: bold; font-size: 13px;
            }}
        """)
        excel_btn.clicked.connect(self._export_excel)
        btn_row.addWidget(excel_btn)

        docx_btn = QPushButton("Export as Word (.docx)")
        docx_btn.setAccessibleName("Export as Word document")
        docx_btn.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
        docx_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        docx_btn.setFixedHeight(APP_SETTINGS["touch_target_min"])
        docx_btn.setStyleSheet(f"""
            QPushButton {{
                background: {c['tertiary']}; color: white;
                border: none; border-radius: 8px;
                padding: 0 24px; font-weight: bold; font-size: 13px;
            }}
        """)
        docx_btn.clicked.connect(self._export_docx)
        btn_row.addWidget(docx_btn)

        btn_row.addStretch()
        content_layout.addLayout(btn_row)

        layout.addWidget(self._content)

    # ------------------------------------------------------------------ data

    def _build_twin_dict(self) -> dict | None:
        """Build a complete dict of all student twin data."""
        if not self._profile:
            return None

        session = self.db.get_session()
        try:
            supports = session.query(SupportEntry).filter(
                SupportEntry.profile_id == self._profile.id
            ).all()

            logs = session.query(TrackingLog).filter(
                TrackingLog.profile_id == self._profile.id
            ).order_by(TrackingLog.created_at.desc()).all()

            insights = session.query(InsightLog).filter(
                InsightLog.profile_id == self._profile.id,
                InsightLog.role == "student",
            ).order_by(InsightLog.created_at.desc()).all()

            twin = {
                "version": "1.0",
                "exported_at": datetime.now(timezone.utc).isoformat(),
                "profile": {
                    "name": self._profile.name,
                    "strengths": self._profile.strengths,
                    "supports_summary": self._profile.supports,
                    "history": self._profile.history,
                    "hopes": self._profile.hopes,
                    "stakeholders": self._profile.stakeholders,
                },
                "support_entries": [
                    {
                        "id": s.id,
                        "category": s.category,
                        "subcategory": s.subcategory,
                        "description": s.description,
                        "udl_mapping": json.loads(s.udl_mapping or "{}"),
                        "pour_mapping": json.loads(s.pour_mapping or "{}"),
                        "status": s.status,
                        "effectiveness_rating": s.effectiveness_rating,
                        "created_at": s.created_at.isoformat() if s.created_at else None,
                    }
                    for s in supports
                ],
                "tracking_logs": [
                    {
                        "id": lg.id,
                        "logged_by_role": lg.logged_by_role,
                        "support_id": lg.support_id,
                        "implementation_notes": lg.implementation_notes,
                        "outcome_notes": lg.outcome_notes,
                        "created_at": lg.created_at.isoformat() if lg.created_at else None,
                    }
                    for lg in logs
                ],
                "insights": [
                    {
                        "id": i.id,
                        "content": i.content,
                        "conversation": i.conversation,
                        "created_at": i.created_at.isoformat() if i.created_at else None,
                    }
                    for i in insights
                ],
            }
            return twin
        finally:
            session.close()

    def _default_filename(self, ext: str) -> str:
        name = self._profile.name.replace(" ", "_") if self._profile else "twin"
        ts = datetime.now().strftime("%Y%m%d_%H%M")
        return f"{name}_twin_{ts}.{ext}"

    # ------------------------------------------------------------------ JSON

    def _export_json(self):
        twin = self._build_twin_dict()
        if not twin:
            QMessageBox.warning(self, "No Data", "No profile data to export.")
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "Save as JSON",
            self._default_filename("json"),
            "JSON Files (*.json)",
        )
        if not path:
            return

        try:
            with open(path, "w") as f:
                json.dump(twin, f, indent=2, ensure_ascii=False)
            QMessageBox.information(
                self, "Exported", f"Twin data saved to:\n{path}"
            )
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Export failed: {e}")

    # ------------------------------------------------------------------ Excel

    def _export_excel(self):
        twin = self._build_twin_dict()
        if not twin:
            QMessageBox.warning(self, "No Data", "No profile data to export.")
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "Save as Excel",
            self._default_filename("xlsx"),
            "Excel Files (*.xlsx)",
        )
        if not path:
            return

        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill

            wb = Workbook()
            header_font = Font(bold=True, size=12)
            header_fill = PatternFill(
                start_color="6F2FA6", end_color="6F2FA6", fill_type="solid"
            )
            header_font_white = Font(bold=True, size=11, color="FFFFFF")

            def style_header_row(ws):
                for cell in ws[1]:
                    cell.font = header_font_white
                    cell.fill = header_fill

            # -- Profile sheet --
            ws = wb.active
            ws.title = "Profile"
            ws.append(["Field", "Value"])
            style_header_row(ws)

            profile = twin["profile"]
            ws.append(["Name", profile["name"]])

            def items_to_text(items):
                if not items:
                    return ""
                texts = []
                for item in items:
                    if isinstance(item, dict):
                        texts.append(item.get("text", str(item)))
                    else:
                        texts.append(str(item))
                return "; ".join(texts)

            ws.append(["Strengths", items_to_text(profile["strengths"])])
            ws.append(["Supports Summary", items_to_text(profile["supports_summary"])])
            ws.append(["History", items_to_text(profile["history"])])
            ws.append(["Hopes & Goals", items_to_text(profile["hopes"])])
            ws.append(["Stakeholders", items_to_text(profile["stakeholders"])])
            ws.column_dimensions["A"].width = 20
            ws.column_dimensions["B"].width = 80

            # -- Support Entries sheet --
            ws2 = wb.create_sheet("Support Entries")
            ws2.append([
                "Category", "Subcategory", "Description", "Status",
                "Effectiveness", "UDL Mapping", "POUR Mapping", "Created"
            ])
            style_header_row(ws2)
            for s in twin["support_entries"]:
                ws2.append([
                    s["category"], s["subcategory"] or "",
                    s["description"], s["status"],
                    s["effectiveness_rating"] or "",
                    json.dumps(s["udl_mapping"]) if s["udl_mapping"] else "",
                    json.dumps(s["pour_mapping"]) if s["pour_mapping"] else "",
                    s["created_at"] or "",
                ])
            for col_letter in ["A", "B", "C", "D", "E", "F", "G", "H"]:
                ws2.column_dimensions[col_letter].width = 20
            ws2.column_dimensions["C"].width = 50

            # -- Tracking Logs sheet --
            ws3 = wb.create_sheet("Tracking Logs")
            ws3.append([
                "Date", "Logged By", "Support ID",
                "Implementation Notes", "Outcome Notes"
            ])
            style_header_row(ws3)
            for lg in twin["tracking_logs"]:
                ws3.append([
                    lg["created_at"] or "",
                    lg["logged_by_role"],
                    lg["support_id"] or "",
                    lg["implementation_notes"] or "",
                    lg["outcome_notes"] or "",
                ])
            ws3.column_dimensions["A"].width = 22
            ws3.column_dimensions["D"].width = 50
            ws3.column_dimensions["E"].width = 50

            # -- Insights sheet --
            ws4 = wb.create_sheet("AI Insights")
            ws4.append(["Date", "Insight Content", "Chat Messages"])
            style_header_row(ws4)
            for ins in twin["insights"]:
                chat_text = ""
                if ins["conversation"]:
                    parts = []
                    for msg in ins["conversation"]:
                        role = "You" if msg.get("role") == "user" else "Assistant"
                        parts.append(f"{role}: {msg.get('content', '')}")
                    chat_text = "\n".join(parts)
                ws4.append([
                    ins["created_at"] or "",
                    ins["content"],
                    chat_text,
                ])
            ws4.column_dimensions["A"].width = 22
            ws4.column_dimensions["B"].width = 80
            ws4.column_dimensions["C"].width = 60

            wb.save(path)
            QMessageBox.information(
                self, "Exported", f"Twin data saved to:\n{path}"
            )
        except ImportError:
            QMessageBox.critical(
                self, "Missing Library",
                "The openpyxl library is required for Excel export.\n"
                "Install it with: pip install openpyxl"
            )
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Export failed: {e}")

    # ------------------------------------------------------------------ DOCX

    def _export_docx(self):
        twin = self._build_twin_dict()
        if not twin:
            QMessageBox.warning(self, "No Data", "No profile data to export.")
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "Save as Word Document",
            self._default_filename("docx"),
            "Word Documents (*.docx)",
        )
        if not path:
            return

        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.table import WD_TABLE_ALIGNMENT

            doc = Document()

            # -- Title --
            title = doc.add_heading("Accessibility Twin Export", level=0)
            title.runs[0].font.color.rgb = RGBColor(0x6F, 0x2F, 0xA6)

            profile = twin["profile"]
            doc.add_paragraph(f"Student: {profile['name']}")
            doc.add_paragraph(
                f"Exported: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
            )
            doc.add_paragraph("")

            # -- Profile Summary --
            doc.add_heading("Profile Summary", level=1)

            def add_item_list(heading, items):
                if not items:
                    return
                doc.add_heading(heading, level=2)
                for item in items:
                    text = item.get("text", str(item)) if isinstance(item, dict) else str(item)
                    priority = item.get("priority", "") if isinstance(item, dict) else ""
                    suffix = f" [{priority}]" if priority and priority != "medium" else ""
                    doc.add_paragraph(f"{text}{suffix}", style="List Bullet")

            add_item_list("Strengths", profile["strengths"])
            add_item_list("Supports Summary", profile["supports_summary"])
            add_item_list("History", profile["history"])
            add_item_list("Hopes & Goals", profile["hopes"])
            add_item_list("Stakeholders", profile["stakeholders"])

            # -- Support Entries --
            if twin["support_entries"]:
                doc.add_heading("Support Entries", level=1)
                table = doc.add_table(
                    rows=1, cols=5, style="Light Grid Accent 1"
                )
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                headers = ["Category", "Description", "Status", "Rating", "Created"]
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                    for p in table.rows[0].cells[i].paragraphs:
                        for run in p.runs:
                            run.font.bold = True

                for s in twin["support_entries"]:
                    row = table.add_row()
                    cat = s["category"]
                    if s["subcategory"]:
                        cat += f" / {s['subcategory']}"
                    row.cells[0].text = cat
                    row.cells[1].text = s["description"]
                    row.cells[2].text = s["status"]
                    row.cells[3].text = (
                        f"{s['effectiveness_rating']}/5"
                        if s["effectiveness_rating"] else "Not rated"
                    )
                    row.cells[4].text = (
                        s["created_at"][:10] if s["created_at"] else ""
                    )

            # -- Tracking Logs --
            if twin["tracking_logs"]:
                doc.add_heading("Tracking Logs", level=1)
                for lg in twin["tracking_logs"]:
                    ts = lg["created_at"][:16] if lg["created_at"] else "Unknown date"
                    doc.add_heading(
                        f"{ts} ({lg['logged_by_role']})", level=3
                    )
                    if lg["implementation_notes"]:
                        p = doc.add_paragraph("")
                        p.add_run("Implementation: ").bold = True
                        p.add_run(lg["implementation_notes"])
                    if lg["outcome_notes"]:
                        p = doc.add_paragraph("")
                        p.add_run("Outcome: ").bold = True
                        p.add_run(lg["outcome_notes"])

            # -- AI Insights --
            if twin["insights"]:
                doc.add_heading("AI Insights", level=1)
                for ins in twin["insights"]:
                    ts = ins["created_at"][:16] if ins["created_at"] else "Unknown"
                    doc.add_heading(f"Insight — {ts}", level=2)
                    for line in ins["content"].split("\n"):
                        if line.strip():
                            doc.add_paragraph(line.strip())

                    # Chat conversation
                    if ins["conversation"]:
                        doc.add_heading("Follow-up Chat", level=3)
                        for msg in ins["conversation"]:
                            role = "You" if msg.get("role") == "user" else "Assistant"
                            p = doc.add_paragraph("")
                            p.add_run(f"{role}: ").bold = True
                            p.add_run(msg.get("content", ""))

            doc.save(path)
            QMessageBox.information(
                self, "Exported", f"Twin data saved to:\n{path}"
            )
        except ImportError:
            QMessageBox.critical(
                self, "Missing Library",
                "The python-docx library is required for Word export.\n"
                "Install it with: pip install python-docx"
            )
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Export failed: {e}")

    # ------------------------------------------------------------------ refresh

    def refresh_data(self):
        user = self.auth.get_current_user()
        if not user:
            return

        session = self.db.get_session()
        try:
            self._profile = session.query(StudentProfile).filter(
                StudentProfile.user_id == user.id
            ).first()
        finally:
            session.close()

        if not self._profile:
            self._empty.setVisible(True)
            self._content.setVisible(False)
            return

        self._empty.setVisible(False)
        self._content.setVisible(True)

        twin = self._build_twin_dict()
        if twin:
            supports_n = len(twin["support_entries"])
            logs_n = len(twin["tracking_logs"])
            insights_n = len(twin["insights"])
            self._summary_label.setText(
                f"Your export includes: {supports_n} support(s), "
                f"{logs_n} tracking log(s), {insights_n} AI insight(s)."
            )
            self._preview.setPlainText(json.dumps(twin, indent=2))
